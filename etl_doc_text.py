#!/usr/bin/env python3
"""
HUJJAT MATNI ETL  (P0-2)
========================
Tenderga biriktirilgan fayllarni MANBADAN yuklab oladi, matnini DETERMINISTIK
parserlar bilan ajratib oladi va `tender_document_text` ga yozadi.

NEGA KERAK: `tender_document` da faqat metama'lumot bor (nom, hajm, tur).
Tenderning haqiqiy mazmuni — texnik topshiriq, talablar, shartlar — PDF/DOCX
ichida qulflangan. Matn ajratilmaguncha na qidiruv, na talab ajratish ishlaydi.

MUHIM: bu skript AI/model CHAQIRMAYDI. Faqat parserlar:
    pdf  -> pypdf            docx -> python-docx
    xlsx -> openpyxl         txt/csv/html/xml -> stdlib
Qolgan formatlar (rar, zip, doc, ...) -> 'unsupported'.

FAYLNI QAYERDAN OLAMIZ (mantiq `api/main.py` dan nusxa olingan — u fayl
o'zgartirilmagan):
    xt-xarid : GET  https://api.xt-xarid.uz/file/<file_id>
    uzex     : POST https://apietender.uzex.uz/api/common/DownloadFile?path=...
               (GET -> 405, brauzer User-Agent shart)
Fayl DISKKA saqlanmaydi — oqim bilan olinadi, matni ajratiladi, o'zi tashlanadi.

STATUSLAR (TZ talabi: "o'qib bo'lmaydigan fayllar 'qo'lda tekshirish talab
etiladi' deb belgilanadi" — 'ok' dan boshqa hammasi shu toifaga kiradi):
    ok | unreadable | unsupported | download_failed | too_large

Ishga tushirish:
    python etl_doc_text.py                      # ishlanmagan barcha hujjatlar
    python etl_doc_text.py --limit 20           # sinov uchun
    python etl_doc_text.py --tender-id 1493304  # bitta tender
    python etl_doc_text.py --force              # qayta ishlash (kesh bekor)
    python etl_doc_text.py --limit 5 --dry-run  # DBga yozmaydi
"""
import argparse
import csv
import io
import logging
import os
import re
import sys
import time
import zipfile
from html.parser import HTMLParser
from typing import Any, Dict, List, Optional, Tuple

import requests

try:
    from dotenv import load_dotenv
except ImportError:  # dotenv bo'lmasa ham muhit o'zgaruvchisidan ishlaydi
    load_dotenv = None

try:
    import psycopg2
    from psycopg2.extras import RealDictCursor
except ImportError:
    psycopg2 = None

# --- Parserlar. Yo'q bo'lsa skript ishlayveradi, o'sha format 'unsupported' ---
try:
    from pypdf import PdfReader
    # pypdf nostandart PDF haqida HAR SAHIFADA ogohlantirish yozadi va
    # natija satrlarini bosib ketadi. Muammo baribir `status` da aks etadi.
    logging.getLogger("pypdf").setLevel(logging.ERROR)
except ImportError:
    PdfReader = None
try:
    import docx as python_docx
except ImportError:
    python_docx = None
try:
    import openpyxl
except ImportError:
    openpyxl = None


# ---------------------------------------------------------------------------
# Sozlamalar
# ---------------------------------------------------------------------------
REQUEST_DELAY = 0.6          # manbaga bosim qilmaslik uchun (soniya)
TIMEOUT       = 90           # yuklab olish (katta PDF'lar bor)
MAX_RETRIES   = 3
RETRY_BACKOFF = 2.0
CHUNK         = 65536        # oqim bo'lagi

MAX_BYTES     = 25 * 1024 * 1024    # 25 MB dan katta -> too_large
MAX_CHARS     = 400_000             # bazaga yoziladigan matn chegarasi
MIN_CHARS     = 20                  # shundan kam matn = amalda bo'sh (skan)
# Chizma/skan PDF'lardan ba'zan bir necha belgi "sizib" chiqadi:
#   '№ № № ³ Ø296 Ø296Ø83Ø88937'  — uzunligi MIN_CHARS dan katta, lekin bu
# MATN emas. Shuning uchun HARF sonini alohida talab qilamiz.
MIN_LETTERS   = 40
PDF_MAX_PAGES = 300                 # 1000 sahifali kitoblar bor — vaqtni cheklaymiz
XLSX_MAX_ROWS = 5000                # varaqdagi qatorlar chegarasi

# `api/main.py` dagi bilan bir xil (o'sha fayl o'zgartirilmadi)
_FILE_URL      = {"xt-xarid": "https://api.xt-xarid.uz/file/{file_id}"}
_POST_DOWNLOAD = {"uzex": "https://apietender.uzex.uz/api/common/DownloadFile"}
_BROWSER_UA    = ("Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                  "AppleWebKit/537.36 (KHTML, like Gecko) "
                  "Chrome/120.0.0.0 Safari/537.36")

# Qaysi kengaytmani qaysi parser oladi
EXT_PDF   = {"pdf"}
EXT_DOCX  = {"docx", "docm"}
EXT_XLSX  = {"xlsx", "xlsm"}
EXT_PLAIN = {"txt", "csv", "htm", "html", "xml", "json", "md"}
# Bilib turib rad etamiz (arxiv / eski binar format / rasm)
EXT_KNOWN_UNSUPPORTED = {"rar", "zip", "7z", "tar", "gz", "doc", "xls", "ppt",
                         "pptx", "jpg", "jpeg", "png", "gif", "tif", "tiff",
                         "bmp", "exe", "dwg", "sig", "p7s", "pfx", "ofd"}

_WS_RE = re.compile(r"[ \t ]+")
_NL_RE = re.compile(r"\n{3,}")


# ---------------------------------------------------------------------------
# Yuklab olish
# ---------------------------------------------------------------------------
def download(session: requests.Session, row: dict) -> Tuple[Optional[bytes], Optional[str]]:
    """Faylni manbadan oladi. -> (baytlar, xato).

    Oqim bilan o'qiladi va MAX_BYTES da UZILADI — server Content-Length
    bermasa ham katta fayl xotirani to'ldirmaydi.
    """
    platform = row.get("source_platform") or "xt-xarid"

    if platform == "xt-xarid" and row.get("file_id"):
        method, url, params = "GET", _FILE_URL["xt-xarid"].format(file_id=row["file_id"]), None
    elif platform in _POST_DOWNLOAD and row.get("file_path"):
        method, url, params = "POST", _POST_DOWNLOAD[platform], {"path": row["file_path"]}
    else:
        return None, f"'{platform}' uchun yuklab olish manzili yo'q"

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            resp = session.request(method, url, params=params,
                                   headers={"User-Agent": _BROWSER_UA},
                                   stream=True, timeout=TIMEOUT)
            resp.raise_for_status()

            # Sarlavhada hajm bo'lsa — yuklamasdan turib rad etamiz
            clen = resp.headers.get("Content-Length")
            if clen and clen.isdigit() and int(clen) > MAX_BYTES:
                resp.close()
                return None, f"__TOO_LARGE__{clen}"

            buf = io.BytesIO()
            size = 0
            for part in resp.iter_content(chunk_size=CHUNK):
                if not part:
                    continue
                size += len(part)
                if size > MAX_BYTES:
                    resp.close()
                    return None, f"__TOO_LARGE__{size}+"
                buf.write(part)
            resp.close()
            return buf.getvalue(), None
        except requests.RequestException as e:
            if attempt == MAX_RETRIES:
                return None, str(e)[:400]
            time.sleep(RETRY_BACKOFF ** attempt)
    return None, "noma'lum xato"


# ---------------------------------------------------------------------------
# Matn ajratish — har biri (matn, sahifa_soni, extractor, xato) qaytaradi
# ---------------------------------------------------------------------------
def clean(text: str) -> str:
    """Ortiqcha bo'shliq/qator — bazada joy va keyingi tahlilda shovqin."""
    # NUL bayt — PostgreSQL TEXT uni QABUL QILMAYDI (ValueError bilan yiqiladi).
    # Buzilgan PDF/DOCX dan chiqib qolishi mumkin, shuning uchun birinchi olib
    # tashlaymiz. Boshqa boshqaruv belgilari ham matnda keraksiz.
    text = text.replace("\x00", " ")
    text = "".join(ch if ch >= " " or ch in "\n\r\t" else " " for ch in text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WS_RE.sub(" ", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    return _NL_RE.sub("\n\n", text).strip()


def extract_pdf(data: bytes) -> Tuple[str, Optional[int], str, Optional[str]]:
    if PdfReader is None:
        return "", None, "pypdf", "pypdf o'rnatilmagan"
    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            # Bo'sh parol ko'p hollarda ochadi (himoya faqat tahrirdan)
            try:
                reader.decrypt("")
            except Exception:  # noqa: BLE001
                return "", None, "pypdf", "PDF parol bilan himoyalangan"
        pages = reader.pages
        total = len(pages)
        parts: List[str] = []
        for i, page in enumerate(pages):
            if i >= PDF_MAX_PAGES:
                break
            try:
                parts.append(page.extract_text() or "")
            except Exception:  # noqa: BLE001
                continue    # bitta sahifa buzilgan bo'lsa qolganini olamiz
            if sum(len(p) for p in parts) > MAX_CHARS:
                break
        return clean("\n".join(parts)), total, "pypdf", None
    except Exception as e:  # noqa: BLE001
        return "", None, "pypdf", str(e)[:400]


def extract_docx(data: bytes) -> Tuple[str, Optional[int], str, Optional[str]]:
    if python_docx is None:
        return "", None, "python-docx", "python-docx o'rnatilmagan"
    try:
        doc = python_docx.Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        n_para = len(doc.paragraphs)
        # JADVALLAR — texnik topshiriqda talablar ko'pincha aynan jadvalda
        for table in doc.tables:
            for r in table.rows:
                cells = [c.text.strip() for c in r.cells]
                line = " | ".join(x for x in cells if x)
                if line:
                    parts.append(line)
        return clean("\n".join(parts)), n_para, "python-docx", None
    except (zipfile.BadZipFile, KeyError) as e:
        # .docx nomi bilan aslida .doc/.rtf bo'lishi mumkin
        return "", None, "python-docx", f"DOCX emas yoki buzilgan: {e}"[:400]
    except Exception as e:  # noqa: BLE001
        return "", None, "python-docx", str(e)[:400]


def extract_xlsx(data: bytes) -> Tuple[str, Optional[int], str, Optional[str]]:
    if openpyxl is None:
        return "", None, "openpyxl", "openpyxl o'rnatilmagan"
    wb = None
    try:
        # read_only + values_only — formula/uslub o'qilmaydi, xotira tejaladi
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: List[str] = []
        for ws in wb.worksheets:
            parts.append(f"[{ws.title}]")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= XLSX_MAX_ROWS:
                    break
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
            if sum(len(p) for p in parts) > MAX_CHARS:
                break
        return clean("\n".join(parts)), len(wb.worksheets), "openpyxl", None
    except (zipfile.BadZipFile, KeyError) as e:
        return "", None, "openpyxl", f"XLSX emas yoki buzilgan: {e}"[:400]
    except Exception as e:  # noqa: BLE001
        return "", None, "openpyxl", str(e)[:400]
    finally:
        if wb is not None:
            try:
                wb.close()
            except Exception:  # noqa: BLE001
                pass


class _TextHTML(HTMLParser):
    """HTML dan ko'rinadigan matnni oladi (script/style tashlanadi)."""
    _SKIP = {"script", "style", "head", "meta", "link"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: List[str] = []
        self._skip = 0

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP:
            self._skip += 1

    def handle_endtag(self, tag):
        if tag in self._SKIP and self._skip:
            self._skip -= 1
        elif tag in ("p", "br", "div", "tr", "li"):
            self.parts.append("\n")

    def handle_data(self, data):
        if not self._skip and data.strip():
            self.parts.append(data)


def extract_plain(data: bytes, ext: str) -> Tuple[str, Optional[int], str, Optional[str]]:
    """txt/csv/html/xml — stdlib. Kodlash noma'lum, ketma-ket sinaymiz."""
    text = None
    for enc in ("utf-8", "utf-8-sig", "cp1251", "cp1252"):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        text = data.decode("utf-8", errors="replace")

    if ext in ("htm", "html", "xml"):
        p = _TextHTML()
        try:
            p.feed(text)
            text = " ".join(p.parts)
        except Exception:  # noqa: BLE001
            text = re.sub(r"<[^>]+>", " ", text)   # zaxira: teglarni olib tashlash
    elif ext == "csv":
        try:
            rows = list(csv.reader(io.StringIO(text)))
            text = "\n".join(" | ".join(c.strip() for c in r if c.strip())
                             for r in rows[:XLSX_MAX_ROWS])
        except Exception:  # noqa: BLE001
            pass

    return clean(text), None, "plain", None


def sniff_ext(row: dict) -> str:
    """Kengaytmani aniqlaydi: file_type -> nom -> content_type."""
    ext = (row.get("file_type") or "").strip().lower().lstrip(".")
    if not ext:
        name = row.get("name") or ""
        if "." in name:
            ext = name.rsplit(".", 1)[-1].strip().lower()
    if not ext:
        ct = (row.get("content_type") or "").lower()
        if "pdf" in ct:
            ext = "pdf"
        elif "wordprocessingml" in ct:
            ext = "docx"
        elif "spreadsheetml" in ct:
            ext = "xlsx"
        elif "text/" in ct:
            ext = "txt"
    return ext


def extract(data: bytes, ext: str) -> Tuple[str, Optional[int], str, Optional[str]]:
    if ext in EXT_PDF:
        return extract_pdf(data)
    if ext in EXT_DOCX:
        return extract_docx(data)
    if ext in EXT_XLSX:
        return extract_xlsx(data)
    if ext in EXT_PLAIN:
        return extract_plain(data, ext)
    return "", None, "", None      # bu yerga kelmasligi kerak


def is_supported(ext: str) -> bool:
    return ext in EXT_PDF or ext in EXT_DOCX or ext in EXT_XLSX or ext in EXT_PLAIN


# ---------------------------------------------------------------------------
# Bitta hujjatni qayta ishlash
# ---------------------------------------------------------------------------
def process(session: requests.Session, row: dict) -> dict:
    """-> tender_document_text ga yoziladigan qator (har doim status bilan)."""
    out = {"tender_id": row["tender_id"], "file_ref": row["file_ref"],
           "text": None, "status": "unreadable", "char_count": None,
           "page_count": None, "error": None, "extractor": None}

    ext = sniff_ext(row)

    # 1. Format — yuklab olishdan OLDIN tekshiramiz (trafik tejaladi)
    if not is_supported(ext):
        out["status"] = "unsupported"
        if not ext:
            out["error"] = "format aniqlanmadi (kengaytma ham, content_type ham yo'q)"
        elif ext in EXT_KNOWN_UNSUPPORTED:
            out["error"] = f"'{ext}' formati qo'llab-quvvatlanmaydi (arxiv/eski binar/rasm)"
        else:
            out["error"] = f"'{ext}' formati uchun parser yo'q"
        return out

    # 2. Hajm — bazadagi metama'lumot bo'yicha
    size = row.get("size_bytes")
    if size and int(size) > MAX_BYTES:
        out["status"] = "too_large"
        out["error"] = f"{int(size) / 1048576:.1f} MB > {MAX_BYTES / 1048576:.0f} MB"
        return out

    # 3. Yuklab olish
    data, err = download(session, row)
    if data is None:
        if err and err.startswith("__TOO_LARGE__"):
            out["status"] = "too_large"
            out["error"] = f"oqimda chegaradan oshdi ({err[13:]} bayt)"
        else:
            out["status"] = "download_failed"
            out["error"] = err
        return out
    if not data:
        out["status"] = "unreadable"
        out["error"] = "bo'sh fayl (0 bayt)"
        return out

    # 4. Matn ajratish
    text, pages, extractor, perr = extract(data, ext)
    out["extractor"] = extractor or None
    out["page_count"] = pages

    if perr:
        out["status"] = "unreadable"
        out["error"] = perr
        return out
    letters = sum(1 for ch in text if ch.isalpha())
    if len(text) < MIN_CHARS or letters < MIN_LETTERS:
        # Eng ko'p uchraydigan holat: SKAN qilingan PDF yoki CHIZMA — OCR kerak.
        out["status"] = "unreadable"
        out["error"] = (f"matn topilmadi (skan/chizma bo'lishi mumkin, OCR kerak; "
                        f"{len(text)} belgi, {letters} harf)"
                        if ext == "pdf" else
                        f"matn topilmadi (bo'sh hujjat; {len(text)} belgi)")
        return out

    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n\n[...matn chegarada qisqartirildi]"
    out["text"] = text
    out["char_count"] = len(text)
    out["status"] = "ok"
    return out


# ---------------------------------------------------------------------------
# DB qatlami
# ---------------------------------------------------------------------------
COLS = ["tender_id", "file_ref", "text", "status", "char_count",
        "page_count", "error", "extractor"]


def fetch_targets(conn, args) -> List[dict]:
    """Qayta ishlanadigan hujjatlar. --force bo'lmasa allaqachon
    ishlanganlari CHIQARIB TASHLANADI — takroriy yuklab olish yo'q."""
    where = ["1=1"]
    params: Dict[str, Any] = {}
    if args.tender_id:
        where.append("d.tender_id = %(tid)s")
        params["tid"] = args.tender_id
    if args.platform:
        where.append("d.source_platform = %(plat)s")
        params["plat"] = args.platform
    if args.file_type:
        where.append("lower(d.file_type) = %(ft)s")
        params["ft"] = args.file_type.strip().lower().lstrip(".")
    if not args.force:
        where.append("t.file_ref IS NULL")

    sql = f"""
        SELECT d.tender_id, d.file_ref, d.file_id, d.file_path, d.name,
               d.size_bytes, d.content_type, d.file_type, d.source_platform
        FROM tender_document d
        LEFT JOIN tender_document_text t
               ON t.tender_id = d.tender_id AND t.file_ref = d.file_ref
        WHERE {' AND '.join(where)}
        ORDER BY d.tender_id DESC, d.file_ref
    """
    if args.limit:
        sql += f" LIMIT {int(args.limit)}"

    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute(sql, params)
        return [dict(r) for r in cur.fetchall()]


def save(conn, rec: dict) -> None:
    """Bitta yozuv (idempotent). Har hujjatdan keyin commit — uzun yurish
    yarmida uzilsa ham ishlangan qism saqlanib qoladi."""
    with conn.cursor() as cur:
        cur.execute(
            f"""INSERT INTO tender_document_text ({','.join(COLS)})
                VALUES ({','.join('%s' for _ in COLS)})
                ON CONFLICT (tender_id, file_ref) DO UPDATE SET
                {','.join(f'{c}=EXCLUDED.{c}' for c in COLS if c not in ('tender_id', 'file_ref'))},
                extracted_at = now()""",
            [rec[c] for c in COLS])
    conn.commit()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main() -> None:
    if load_dotenv:
        load_dotenv()

    ap = argparse.ArgumentParser(description="Tender hujjatlaridan matn ajratish ETL")
    ap.add_argument("--limit", type=int, help="Nechta hujjat (sinov uchun)")
    ap.add_argument("--tender-id", type=int, help="Faqat shu tender hujjatlari")
    ap.add_argument("--platform", help="Faqat shu manba ('xt-xarid' | 'uzex')")
    ap.add_argument("--file-type", help="Faqat shu tur ('pdf' | 'docx' | 'xlsx' ...)")
    ap.add_argument("--force", action="store_true",
                    help="Allaqachon ishlanganlarni ham qayta yuklab oladi")
    ap.add_argument("--dry-run", action="store_true", help="DBga yozmaydi")
    ap.add_argument("--quiet", action="store_true", help="Har fayl uchun satr chiqarmaydi")
    ap.add_argument("--dsn", default=os.environ.get("XT_DB_DSN"))
    args = ap.parse_args()

    if not args.dsn:
        sys.exit("XATO: DSN yo'q. --dsn yoki XT_DB_DSN o'rnating.")
    if psycopg2 is None:
        sys.exit("XATO: pip install psycopg2-binary")

    conn = psycopg2.connect(args.dsn)
    rows = fetch_targets(conn, args)

    if not rows:
        print("Qayta ishlanadigan hujjat yo'q "
              "(hammasi ishlangan bo'lishi mumkin — --force bilan qayta yurgizing).")
        conn.close()
        return

    print(f"[1/2] {len(rows)} ta hujjat qayta ishlanadi"
          f"{' (--force)' if args.force else ''}...")
    print(f"      Taxminiy vaqt: ~{len(rows) * (REQUEST_DELAY + 1.2) / 60:.1f} daqiqa\n")

    session = requests.Session()
    counts: Dict[str, int] = {}
    total_chars = 0
    downloaded = 0

    for i, row in enumerate(rows, 1):
        rec = process(session, row)
        counts[rec["status"]] = counts.get(rec["status"], 0) + 1
        total_chars += rec["char_count"] or 0

        if not args.quiet:
            name = (row.get("name") or row["file_ref"])[:48]
            extra = (f"{rec['char_count']} belgi" if rec["status"] == "ok"
                     else (rec["error"] or "")[:60])
            print(f"  [{i}/{len(rows)}] #{row['tender_id']} {name} "
                  f"-> {rec['status']} ({extra})")

        if not args.dry_run:
            try:
                save(conn, rec)
            except Exception as e:  # noqa: BLE001
                conn.rollback()
                print(f"    ! DB xato: {e}", file=sys.stderr)

        # Kechikish FAQAT tarmoqqa chiqqan bo'lsak (unsupported/too_large tez o'tadi)
        if rec["status"] not in ("unsupported", "too_large"):
            downloaded += 1
            time.sleep(REQUEST_DELAY)

    conn.close()

    print(f"\n[2/2] Tayyor. Yuklab olingan fayl: {downloaded}, "
          f"jami matn: {total_chars:,} belgi")
    for st in sorted(counts, key=lambda k: -counts[k]):
        print(f"      {st:<16} {counts[st]}")
    manual = sum(v for k, v in counts.items() if k != "ok")
    print(f"      -> qo'lda tekshirish talab etiladi: {manual}")
    if args.dry_run:
        print("      (--dry-run — DBga yozilmadi)")


if __name__ == "__main__":
    main()
